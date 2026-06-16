#!/usr/bin/env python3
"""Export a question set to a Word (.docx) file from the command line.

Usage:
    python make_docx.py <questions.json> [output.docx]

This mirrors the in-browser "Export Word" button: questions grouped by topic,
then category / conceptual level, options shuffled with the balanced algorithm,
and the correct answer printed at the END of each MCQ (`Answer: B`) with the
optional explanation. No correct option is highlighted inline.

Requires python-docx (`pip install python-docx`).
"""
from __future__ import annotations

import os
import sys

import mcqlib

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    raise SystemExit("python-docx is required: pip install python-docx")


def main(qpath, out=None):
    data = mcqlib.load(qpath)
    lp = mcqlib.level_points(data)
    doc = Document()

    doc.add_heading(data.get("title", "Quiz"), level=0)
    if data.get("source"):
        doc.add_paragraph(f"Source: {data['source']}")

    qnum = 0
    for topic in data["topics"]:
        doc.add_heading(topic["name"], level=1)
        flat = mcqlib.shuffle_set(
            [dict(q, topic=topic["name"]) for q in sorted(topic["questions"], key=mcqlib.sort_key)],
            seed=topic["name"],
        )
        last_group = None
        for q in flat:
            if q["category"] == "conceptual":
                group = f"Conceptual Understanding — {mcqlib.LEVEL_LABEL[q['level']]} ({lp[q['level']]} pt)"
            else:
                group = mcqlib.CATEGORY_LABEL[q["category"]]
            if group != last_group:
                doc.add_heading(group, level=2)
                last_group = group

            qnum += 1
            p = doc.add_paragraph()
            p.add_run(f"Q{qnum}. {q['question']}").bold = True
            for i, opt in enumerate(q["options"]):
                doc.add_paragraph(f"{mcqlib.LETTERS[i]}. {opt}")
            ans = doc.add_paragraph()
            ans.add_run(f"Answer: {q['answer_letter']}").bold = True
            if q.get("explanation"):
                exp = doc.add_paragraph()
                r = exp.add_run(f"Explanation: {q['explanation']}")
                r.italic = True
                r.font.size = Pt(10)

    if out is None:
        base = os.path.splitext(os.path.basename(qpath))[0]
        out = os.path.join(os.path.dirname(os.path.abspath(qpath)), base + ".docx")
    doc.save(out)
    print(f"Wrote {out}  ({qnum} questions)")
    return out


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(2)
    main(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
